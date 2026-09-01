# -*- coding: utf-8 -*-
"""
Формування підсумкового звіту по рейсу (.docx) — водій, машина, вантаж,
хронологія по фазах (Завантаження/Старт/Стоп/Вивантаження), паузи,
додаткові зупинки, ЧП та фото.
"""
import json
import os
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import database as db
from checklists import STEPS, PHASE_ORDER, PHASE_LABELS_PLAIN, steps_of_phase
from config import REPORTS_DIR

BRAND_BLUE = RGBColor(0x1F, 0x3A, 0x5F)
BRAND_GRAY = RGBColor(0x59, 0x59, 0x59)
OK_GREEN = RGBColor(0x1E, 0x7D, 0x32)
WARN_RED = RGBColor(0xB0, 0x00, 0x20)
WARN_ORANGE = RGBColor(0xB2, 0x6A, 0x00)


def _fmt_dt(iso_str):
    if not iso_str:
        return "—"
    try:
        return datetime.fromisoformat(iso_str).strftime("%d.%m.%Y %H:%M")
    except Exception:
        return iso_str


def _shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def build_trip_report(trip_id: int) -> str:
    trip = db.get_trip(trip_id)
    driver = db.get_driver(trip["driver_id"])
    answers = db.get_all_answers(trip_id)
    checked_keys = {a["item_key"] for a in answers if a["checked_at"]}
    photos = db.get_photos(trip_id)
    incidents = db.get_incidents(trip_id)
    pauses = db.get_optional_events(trip_id, "pause")
    extra_stops = db.get_optional_events(trip_id, "extra_stop")

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    # --- Заголовок ---
    title = doc.add_paragraph()
    run = title.add_run("TRANCOM  ·  Звіт по рейсу")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = BRAND_BLUE

    subtitle = doc.add_paragraph()
    sub_run = subtitle.add_run(f"Рейс №{trip_id}")
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = BRAND_GRAY

    doc.add_paragraph()

    # --- Загальна інформація ---
    info_table = doc.add_table(rows=0, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    info_rows = [
        ("Водій", driver["full_name"] if driver else "—"),
        ("Телефон водія", driver["phone"] if driver and driver["phone"] else "—"),
        ("Тягач", trip["tractor_number"] or "—"),
        ("Причіп/платформа", trip["trailer_number"] or "—"),
        ("Вантаж", trip["cargo_description"] or "—"),
        ("Маршрут", trip["route"] or "—"),
        ("Початок рейсу", _fmt_dt(trip["started_at"])),
        ("Завершення рейсу", _fmt_dt(trip["finished_at"])),
        ("Статус", "Завершено" if trip["status"] == "done" else trip["status"]),
        ("Пауз за рейс", str(len(pauses))),
        ("Додаткових зупинок", str(len(extra_stops))),
    ]
    for label, value in info_rows:
        row_cells = info_table.add_row().cells
        row_cells[0].width = Cm(4.5)
        row_cells[1].width = Cm(11)
        r0 = row_cells[0].paragraphs[0].add_run(label)
        r0.bold = True
        r0.font.size = Pt(10.5)
        _shade_cell(row_cells[0], "EFF3F8")
        r1 = row_cells[1].paragraphs[0].add_run(str(value))
        r1.font.size = Pt(10.5)

    doc.add_paragraph()

    # --- Інциденти нагорі, це найважливіше ---
    if incidents:
        h = doc.add_paragraph()
        hr = h.add_run("⚠ Зафіксовані ЧП")
        hr.bold = True
        hr.font.size = Pt(13)
        hr.font.color.rgb = WARN_RED
        for inc in incidents:
            p = doc.add_paragraph()
            r = p.add_run(f"[{_fmt_dt(inc['created_at'])}] {inc['description']}")
            r.font.size = Pt(10.5)
            r.font.color.rgb = WARN_RED
        doc.add_paragraph()

    # --- Хронологія обов'язкових фаз ---
    h2 = doc.add_paragraph()
    hr2 = h2.add_run("Хронологія рейсу: Завантаження → Старт → Стоп → Вивантаження")
    hr2.bold = True
    hr2.font.size = Pt(14)
    hr2.font.color.rgb = BRAND_BLUE

    for phase in PHASE_ORDER:
        ph = doc.add_paragraph()
        ph_run = ph.add_run(PHASE_LABELS_PLAIN[phase])
        ph_run.bold = True
        ph_run.font.size = Pt(12.5)
        ph_run.font.color.rgb = BRAND_BLUE

        for step in steps_of_phase(phase):
            sp = doc.add_paragraph()
            sp.paragraph_format.left_indent = Cm(0.3)
            sp_run = sp.add_run(step["title"])
            sp_run.bold = True
            sp_run.italic = True
            sp_run.font.size = Pt(10.5)
            sp_run.font.color.rgb = BRAND_GRAY

            for item in step["items"]:
                done = item["key"] in checked_keys
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.7)
                mark = "✔" if done else "✘"
                r = p.add_run(f"{mark}  {item['text']}")
                r.font.size = Pt(10)
                r.font.color.rgb = OK_GREEN if done else WARN_RED

    doc.add_paragraph()

    # --- Паузи ---
    if pauses:
        h = doc.add_paragraph()
        hr = h.add_run(f"Паузи в дорозі ({len(pauses)})")
        hr.bold = True
        hr.font.size = Pt(13)
        hr.font.color.rgb = WARN_ORANGE
        for i, pause in enumerate(pauses, 1):
            p = doc.add_paragraph()
            r = p.add_run(f"Пауза №{i} — {_fmt_dt(pause['created_at'])}")
            r.bold = True
            r.font.size = Pt(10.5)
            try:
                items = json.loads(pause["items_checked_json"] or "[]")
            except Exception:
                items = []
            for text in items:
                ip = doc.add_paragraph()
                ip.paragraph_format.left_indent = Cm(0.5)
                ipr = ip.add_run(f"✔ {text}")
                ipr.font.size = Pt(9.5)
                ipr.font.color.rgb = OK_GREEN
        doc.add_paragraph()

    # --- Додаткові зупинки ---
    if extra_stops:
        h = doc.add_paragraph()
        hr = h.add_run(f"Додаткові зупинки ({len(extra_stops)})")
        hr.bold = True
        hr.font.size = Pt(13)
        hr.font.color.rgb = WARN_ORANGE
        for i, stop in enumerate(extra_stops, 1):
            p = doc.add_paragraph()
            note = stop["note"] or "без опису причини"
            r = p.add_run(f"№{i} — {_fmt_dt(stop['created_at'])}: {note}")
            r.font.size = Pt(10)
        doc.add_paragraph()

    # --- Фотофіксація ---
    if photos:
        h3 = doc.add_paragraph()
        hr3 = h3.add_run("Фотофіксація")
        hr3.bold = True
        hr3.font.size = Pt(14)
        hr3.font.color.rgb = BRAND_BLUE

        kind_labels = {
            "checklist": "Чек-лист",
            "pause": "Пауза",
            "extra_stop": "Додаткова зупинка",
            "incident": "ЧП",
        }
        for ph in photos:
            cap = doc.add_paragraph()
            label = kind_labels.get(ph["kind"], ph["kind"])
            cap_run = cap.add_run(f"{label} · {_fmt_dt(ph['created_at'])}")
            cap_run.font.size = Pt(9)
            cap_run.font.color.rgb = BRAND_GRAY
            if os.path.exists(ph["file_path"]):
                try:
                    doc.add_picture(ph["file_path"], width=Cm(9))
                except Exception:
                    doc.add_paragraph(f"[файл фото недоступний: {ph['file_path']}]")
            doc.add_paragraph()

    # --- Футер ---
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run(
        f"Автоматично згенеровано ботом TRANCOM · {datetime.now().strftime('%d.%m.%Y %H:%M')}"
    )
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = BRAND_GRAY

    out_path = os.path.join(REPORTS_DIR, f"reys_{trip_id}.docx")
    doc.save(out_path)
    return out_path
